import os
import pdfplumber
import docx

def extract_text_from_pdf(file_path):
    """
    Mengambil seluruh teks dari file PDF, halaman per halaman.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File tidak ditemukan: {file_path}")
        
    text_content = []
    metadata = {
        "pages": 0,
        "tables_count": 0,
        "characters": 0,
        "words": 0
    }
    
    with pdfplumber.open(file_path) as pdf:
        metadata["pages"] = len(pdf.pages)
        for i, page in enumerate(pdf.pages):
            # Ambil teks
            page_text = page.extract_text()
            if page_text:
                text_content.append(f"--- Halaman {i+1} ---\n{page_text}")
                metadata["characters"] += len(page_text)
                metadata["words"] += len(page_text.split())
            
            # Cek tabel
            tables = page.extract_tables()
            if tables:
                metadata["tables_count"] += len(tables)
                
    return "\n\n".join(text_content), metadata

def extract_text_from_docx(file_path):
    """
    Mengambil seluruh teks dari file DOCX.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File tidak ditemukan: {file_path}")
        
    doc = docx.Document(file_path)
    text_content = []
    metadata = {
        "pages": "N/A (Word Doc)",  # Word doc tidak memiliki representasi halaman fisik yang tetap tanpa dirender
        "tables_count": len(doc.tables),
        "characters": 0,
        "words": 0
    }
    
    # Ambil teks dari paragraf
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            text_content.append(para.text)
            metadata["characters"] += len(para.text)
            metadata["words"] += len(para.text.split())
            
    # Ambil teks dari tabel
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                combined_row_text = " | ".join(row_text)
                text_content.append(f"[Tabel Baris] {combined_row_text}")
                metadata["characters"] += len(combined_row_text)
                metadata["words"] += len(combined_row_text.split())

    return "\n".join(text_content), metadata

def parse_document(file_path):
    """
    Mendeteksi ekstensi file dan mengekstrak konten serta metadata-nya.
    """
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return extract_text_from_pdf(file_path)
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Format file '{ext}' tidak didukung. Harap masukkan .pdf atau .docx")
